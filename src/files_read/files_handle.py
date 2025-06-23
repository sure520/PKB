import os
import traceback
import re
import logging
import uuid
from typing import Dict, List, Tuple, Union, Optional, Any
from urllib.parse import quote
import docx
from docx import Document
import fitz  # PyMuPDF
from fastapi import HTTPException
import os
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from src.files_read.ocr_extract import extract_text_with_langchain_qwen_ocr as qwen_ocr
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyMuPDFLoader,
    UnstructuredMarkdownLoader,
    TextLoader,
    CSVLoader,
    JSONLoader
)

logger = logging.getLogger(__name__)

"""create by haozi
    2025-03-14
    文件处理类，支持从不同格式的文件中提取内容，返回带章节格式的json内容
"""


class FileHandler:
    """文件处理类，支持从不同格式的文件中提取内容"""

    def __init__(self):
        """初始化文件处理器"""
        pass

    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        处理文件并返回提取的文本内容及文件结构

        参数:
            file_path (str): 文件路径

        返回:
            Dict[str, Any]: 包含文本内
            容和结构信息的字典
        """
        try:
            # 检查文件是否存在和可访问
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                raise ValueError(f"文件为空: {file_path}")

            file_ext = os.path.splitext(file_path)[1].lower()
            file_name = os.path.basename(file_path)

            print(f"正在处理文件: {file_path}, 类型: {file_ext}, 大小: {file_size} 字节")

            result = {
                "file_name": file_name,
                "file_path": file_path,
                "file_type": file_ext.replace('.', ''),
                "content": "",
                "structure": {},
                "paragraphs": []  # 添加段落结构
            }

            try:
                if file_ext == '.txt':
                    result["content"] = self._read_text_file(file_path)
                    result["paragraphs"] = self._split_into_paragraphs(result["content"])
                elif file_ext == '.docx':
                    content, structure = self._read_docx_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                elif file_ext == '.doc':
                    content, structure = self._read_doc_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                elif file_ext == '.pdf':
                    content, structure = self._read_pdf_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                elif file_ext in ['.md', '.markdown']:
                    content, structure = self._read_markdown_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif']:
                    result["content"] = self._read_image_file(file_path)
                    result["paragraphs"] = self._split_into_paragraphs(result["content"])
                elif file_ext in ['.ppt', '.pptx']:
                    content, structure = self._read_ppt_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                else:
                    message = f"不支持的文件格式: {file_ext}"
                    # print(message)
                    result["error"] = message
                    result["content"] = f"无法处理此文件格式: {file_ext}"
            except Exception as e:
                # 捕获并记录特定文件处理错误
                error_message = f"处理文件内容时出错: {str(e)}"
                print(error_message)
                traceback.print_exc()
                result["error"] = error_message
                result["content"] = f"文件处理错误: {str(e)}"

            # 确保始终返回有效内容，即使是错误消息
            if not result["content"]:
                result["content"] = "未能提取文件内容"

            return result

        except Exception as e:
            # 捕获顶层异常
            error_message = f"处理文件时出错: {str(e)}"
            print(error_message)
            traceback.print_exc()
            return {
                "file_name": os.path.basename(file_path),
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower().replace('.', ''),
                "content": f"文件处理失败: {str(e)}",
                "structure": {},
                "paragraphs": [],
                "error": str(e)
            }

    def _split_into_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """
        将文本内容分割成段落结构，确保表格和图片与其上下文保持在一起

        参数:
            text (str): 要分割的文本内容

        返回:
            List[Dict[str, Any]]: 段落列表，每个段落包含内容和位置信息
        """
        paragraphs = []

        # 第一步：检测所有表格和图片位置
        table_pattern = re.compile(r'\n表格\s+\d+:[\s\S]+?(?=\n\n表格\s+\d+:|$)')
        image_pattern = re.compile(r'\n插图开始[\s\S]+?插图结束\n')

        table_matches = list(table_pattern.finditer(text))
        image_matches = list(image_pattern.finditer(text))

        # 收集所有特殊块的位置（表格和图片）
        special_blocks = []
        for m in table_matches:
            special_blocks.append(("table", m.start(), m.end()))
        for m in image_matches:
            special_blocks.append(("image", m.start(), m.end()))

        # 按照在文档中的位置排序
        special_blocks.sort(key=lambda x: x[1])

        # 第二步：查找这些特殊块的上下文段落
        context_blocks = []

        # 确定段落边界
        all_paragraphs = re.split(r'\n\s*\n', text)
        paragraph_positions = []

        pos = 0
        for para in all_paragraphs:
            start = pos
            end = pos + len(para)
            if para.strip():  # 只记录非空段落
                paragraph_positions.append((start, end))
            pos = end + 2  # +2 是为了考虑分隔符 \n\n

        # 为每个特殊块找到其上下文段落
        for block_type, block_start, block_end in special_blocks:
            # 查找块前的相关段落作为上下文
            context_start = 0
            for para_start, para_end in paragraph_positions:
                if para_end < block_start and para_end > context_start:
                    # 检查段落是否与特殊块相关
                    para_text = text[para_start:para_end]

                    # 表格相关：包含"表"字或距离较近
                    if block_type == "table" and ("表" in para_text or (block_start - para_end) < 100):
                        context_start = para_start

                    # 图片相关：包含"图"字或"如下所示"等描述性词语或距离较近
                    elif block_type == "image" and ("图" in para_text or
                                                    "如下所示" in para_text or
                                                    "示意图" in para_text or
                                                    (block_start - para_end) < 100):
                        context_start = para_start

            if context_start > 0:
                context_blocks.append((context_start, block_end, block_type))
            else:
                context_blocks.append((block_start, block_end, block_type))

        # 合并重叠的上下文块（可能一段文字同时描述了表格和图片）
        merged_blocks = []
        if context_blocks:
            # 按起始位置排序
            context_blocks.sort(key=lambda x: x[0])
            current_block = context_blocks[0]

            for next_start, next_end, next_type in context_blocks[1:]:
                curr_start, curr_end, curr_type = current_block

                # 如果有重叠，合并块
                if next_start <= curr_end:
                    current_block = (curr_start, max(curr_end, next_end), f"{curr_type}_{next_type}")
                else:
                    merged_blocks.append(current_block)
                    current_block = (next_start, next_end, next_type)

            merged_blocks.append(current_block)

        # 第三步：处理常规文本和特殊块
        processed_ranges = []
        for start, end, block_type in merged_blocks:
            processed_ranges.append((start, end))

        # 处理未被特殊块覆盖的文本
        pos = 0
        while pos < len(text):
            # 检查当前位置是否已处理
            inside_processed = False
            for start, end in processed_ranges:
                if start <= pos < end:
                    inside_processed = True
                    pos = end
                    break

            if inside_processed:
                continue

            # 找到下一个已处理区域
            next_start = len(text)
            for start, end in processed_ranges:
                if start > pos and start < next_start:
                    next_start = start

            # 处理这段普通文本
            if next_start > pos:
                segment_text = text[pos:next_start].strip()
                if segment_text:
                    regular_paragraphs = self._process_regular_text(segment_text, pos)
                    paragraphs.extend(regular_paragraphs)
                pos = next_start

        # 添加特殊块（表格/图片及其上下文）
        for idx, (start, end, block_type) in enumerate(merged_blocks):
            block_text = text[start:end].strip()
            if block_text:
                # 根据块类型设置不同的标记
                contains_table = "table" in block_type
                contains_image = "image" in block_type

                paragraphs.append({
                    "id": f"special_block_{idx}",
                    "content": block_text,
                    "position": start,
                    "length": end - start,
                    "is_heading": False,
                    "heading_level": 0,
                    "contains_table": contains_table,
                    "contains_image": contains_image
                })

        # 按位置排序段落
        paragraphs.sort(key=lambda x: x["position"])

        return paragraphs

    def _process_regular_text(self, text: str, base_position: int = 0) -> List[Dict[str, Any]]:
        """
        处理普通文本（非表格相关）并分割成段落

        参数:
            text (str): 要处理的文本
            base_position (int): 文本在原始文档中的起始位置

        返回:
            List[Dict[str, Any]]: 处理后的段落列表
        """
        result = []
        raw_paragraphs = re.split(r'\n\s*\n', text)

        start_pos = base_position
        for i, para in enumerate(raw_paragraphs):
            if para.strip():  # 忽略空段落
                # 检测是否为标题
                is_heading = False
                heading_level = 0
                header_match = re.match(r'^(#{1,6})\s+(.+)$', para.strip())
                if header_match:
                    is_heading = True
                    heading_level = len(header_match.group(1))

                paragraph = {
                    "id": f"p_{start_pos}",
                    "content": para.strip(),
                    "position": start_pos,
                    "length": len(para),
                    "is_heading": is_heading,
                    "heading_level": heading_level if is_heading else 0,
                    "contains_table": False,
                    "contains_image": False  # 添加图片标记
                }
                result.append(paragraph)

            start_pos += len(para) + 2  # +2 是为了考虑分隔符 \n\n

        return result

    def _read_text_file(self, file_path: str) -> List[Document]:
        """读取文本文件"""
        return TextLoader(file_path,  encoding="utf-8").load()

    def _read_docx_file(self, file_path: str) -> List[Document]:
        """读取docx文件"""
        document = None
        try:
            document = docx.Document(file_path)
            # 提取段落文本
            content = "\n".join([para.text for para in document.paragraphs if para.text.strip()])

            # 创建结构化信息
            structure = {
                "paragraphs": [para.text for para in document.paragraphs if para.text.strip()],
                "tables": []
            }

            # 提取表格内容
            for table in document.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                structure["tables"].append(table_data)

            return content, structure
        except Exception as e:
            logger.error(f"读取docx文件失败: {str(e)}")
            raise ValueError(f"无法处理docx文件: {str(e)}")

    def _read_doc_file(self, file_path):
        """读取传统的doc文件"""
        try:
            print(f"开始处理.doc文件: {file_path}")

            # 尝试主要方法：使用pywin32（仅在Windows环境下有效）
            try:
                print(f"尝试使用pywin32处理.doc文件...")
                import win32com.client
                import os
                # 创建临时文件路径来保存转换后的docx
                temp_path = file_path + ".docx"

                print(f"初始化Word应用程序...")
                # 使用Word应用程序转换doc为docx
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False

                print(f"打开原始文档: {file_path}")
                # 尝试以只读方式打开文档
                try:
                    doc = word.Documents.Open(file_path, ReadOnly=True)
                except Exception as e:
                    print(f"以只读模式打开失败: {str(e)}，尝试普通模式")
                    doc = word.Documents.Open(file_path)

                print(f"将文档保存为docx格式: {temp_path}")
                # 保存为docx
                doc.SaveAs(temp_path, 16)  # 16表示docx格式
                doc.Close()
                word.Quit()

                print(f"转换成功，开始读取转换后的docx文件...")
                # 读取转换后的docx文件
                content, structure = self._read_docx_file(temp_path)

                # 删除临时文件
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    print(f"已删除临时docx文件")

                return content, structure
            except ImportError as e:
                print(f"pywin32导入失败: {str(e)}")
                print(f"将尝试其他方法...")
                # 如果pywin32不可用，尝试使用其他方法
                pass
            except Exception as e:
                print(f"pywin32处理失败: {str(e)}")
                traceback.print_exc()
                print(f"将尝试其他方法...")
                pass

            # 备选方法1: 使用docx2txt
            try:
                print(f"尝试使用docx2txt处理.doc文件...")
                import docx2txt
                content = docx2txt.process(file_path)
                print(f"docx2txt处理成功，提取内容长度: {len(content)}")
                paragraphs = content.split('\n\n')
                structure = {
                    "paragraphs": paragraphs,
                    "tables": []
                }
                return content, structure
            except ImportError as e:
                print(f"docx2txt导入失败: {str(e)}")
                print(f"将尝试其他方法...")
                pass
            except Exception as e:
                print(f"docx2txt处理失败: {str(e)}")
                traceback.print_exc()
                print(f"将尝试其他方法...")
                pass

            # 备选方法2: 尝试使用pdf2docx
            try:
                print(f"尝试使用pdf2docx处理.doc文件...")
                from pdf2docx import parse
                # 创建临时文件
                temp_pdf = file_path + ".pdf"
                temp_docx = file_path + ".pdf.docx"

                # 尝试先转换为PDF（需要Word）
                try:
                    import win32com.client
                    print(f"使用Word转换doc到pdf...")
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(file_path)
                    doc.SaveAs(temp_pdf, 17)  # 17表示PDF格式
                    doc.Close()
                    word.Quit()

                    if os.path.exists(temp_pdf):
                        print(f"转换为PDF成功，将PDF转换为docx...")
                        parse(temp_pdf, temp_docx)

                        if os.path.exists(temp_docx):
                            print(f"转换为docx成功，读取内容...")
                            content, structure = self._read_docx_file(temp_docx)

                            # 清理临时文件
                            os.remove(temp_pdf)
                            os.remove(temp_docx)

                            return content, structure
                except Exception as e:
                    print(f"pdf2docx转换失败: {str(e)}")
                    traceback.print_exc()

                    # 清理临时文件
                    if os.path.exists(temp_pdf):
                        os.remove(temp_pdf)
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
            except ImportError as e:
                print(f"pdf2docx导入失败: {str(e)}")
                print(f"将尝试其他方法...")
                pass

            # 备选方法3：使用textract
            try:
                print(f"尝试使用textract处理.doc文件...")
                import textract
                content = textract.process(file_path).decode('utf-8')
                print(f"textract处理成功，提取内容长度: {len(content)}")
                paragraphs = content.split('\n\n')
                structure = {
                    "paragraphs": paragraphs,
                    "tables": []
                }
                return content, structure
            except ImportError as e:
                print(f"textract导入失败: {str(e)}")
                print(f"将尝试其他方法...")
                pass
            except Exception as e:
                print(f"textract处理失败: {str(e)}")
                traceback.print_exc()
                print(f"将尝试其他方法...")
                pass

            # 备选方法4：使用antiword（需要系统安装antiword）
            try:
                print(f"尝试使用antiword处理.doc文件...")
                import subprocess
                result = subprocess.run(['antiword', file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                content = result.stdout.decode('utf-8')
                print(f"antiword处理成功，提取内容长度: {len(content)}")
                paragraphs = content.split('\n\n')
                structure = {
                    "paragraphs": paragraphs,
                    "tables": []
                }
                return content, structure
            except (ImportError, FileNotFoundError) as e:
                print(f"antiword导入或执行失败: {str(e)}")
                print(f"将尝试其他方法...")
                pass
            except Exception as e:
                print(f"antiword处理失败: {str(e)}")
                traceback.print_exc()
                print(f"将尝试其他方法...")
                pass

            # 备选方法5：使用python-docx直接打开（有可能失败）
            try:
                print(f"尝试使用python-docx直接打开.doc文件（不太可能成功）...")
                content, structure = self._read_docx_file(file_path)
                print(f"python-docx处理成功，内容长度: {len(content)}")
                return content, structure
            except Exception as e:
                print(f"python-docx直接处理失败: {str(e)}")
                traceback.print_exc()
                pass

            # 备选方法6：使用olefile尝试直接提取文本
            try:
                print(f"尝试使用olefile提取.doc文件中的文本...")
                import olefile
                if olefile.isOleFile(file_path):
                    ole = olefile.OleFile(file_path)
                    if ole.exists('WordDocument'):
                        print(f"找到WordDocument流，尝试提取内容...")
                        # 这里只是简单地提取，实际上需要更复杂的处理来正确解析Word二进制格式
                        stream = ole.openstream('WordDocument')
                        data = stream.read()
                        ole.close()

                        # 简单地提取可能的文本
                        content = ""
                        for i in range(0, len(data), 2):
                            if i + 1 < len(data):
                                if 32 <= data[i] <= 126 and data[i + 1] == 0:  # 基本ASCII文本
                                    content += chr(data[i])

                        if len(content) > 100:
                            print(f"olefile提取到 {len(content)} 字符")
                            paragraphs = content.split('\n\n')
                            structure = {
                                "paragraphs": paragraphs,
                                "tables": []
                            }
                            return content, structure
                        else:
                            print(f"olefile未能提取到足够的文本内容")
                    else:
                        print(f"文件不包含WordDocument流")
                else:
                    print(f"文件不是有效的OLE文件")
            except ImportError as e:
                print(f"olefile导入失败: {str(e)}")
                print(f"将尝试其他方法...")
                pass
            except Exception as e:
                print(f"olefile处理失败: {str(e)}")
                traceback.print_exc()
                print(f"将尝试其他方法...")
                pass

            # 尝试最后手段：以二进制方式读取并尝试提取文本
            try:
                print(f"尝试最后手段：以二进制方式读取并尝试提取文本...")
                with open(file_path, 'rb') as f:
                    data = f.read()
                # 尝试从二进制数据中提取ASCII文本
                content = ""
                for byte in data:
                    if 32 <= byte <= 126:  # 基本ASCII可打印字符
                        content += chr(byte)
                    elif byte in [9, 10, 13]:  # 制表符、换行符、回车符
                        content += chr(byte)

                if len(content) > 100:  # 假设至少有一些有意义的文本
                    print(f"二进制读取提取了 {len(content)} 字符的文本")
                    structure = {"paragraphs": content.split('\n\n'), "tables": []}
                    return content, structure
                else:
                    print("二进制读取未能提取有意义的文本")
            except Exception as e:
                print(f"二进制读取失败: {str(e)}")
                traceback.print_exc()

            error_msg = "无法处理此doc文件，所有处理方法均已失败"
            print(error_msg)
            raise ValueError(error_msg)

        except Exception as e:
            print(f"doc文件处理过程中出现异常: {str(e)}")
            traceback.print_exc()
            raise

    def _read_pdf_file(self, file_path: str) -> Tuple[str, Dict]:
        """使用PyMuPDF提取PDF文本和表格"""
        full_text = ""
        structure = {
            "pages": 0,
            "toc": [],
            "metadata": {},
            "images_count": 0,
            "tables_count": 0
        }

        total_tables = 0
        total_images = 0

        try:
            # 使用PyMuPDF打开PDF
            doc = fitz.open(file_path)
            structure["pages"] = len(doc)

            # 尝试获取目录和元数据（如果出现错误，继续处理）
            try:
                structure["toc"] = doc.get_toc()
                structure["metadata"] = doc.metadata
            except Exception as e:
                print(f"提取PDF元数据时出错: {str(e)}")

            # 逐页处理
            for page_idx, page in enumerate(doc):
                page_text = ""

                # 安全提取页面文本
                try:
                    page_text = page.get_text()
                except Exception as e:
                    print(f"提取页面 {page_idx + 1} 文本时出错: {str(e)}")
                    page_text = f"[无法提取页面 {page_idx + 1} 文本: {str(e)}]"

                # 尝试提取表格
                try:
                    tables = page.find_tables()
                    print(f"页面 {page_idx + 1} 表格查找结果: {tables}")

                    # 检查表格对象类型并打印更多信息
                    if tables:
                        print(f"表格对象类型: {type(tables)}")
                        print(f"表格对象属性: {dir(tables)}")

                    # 尝试多种方式处理表格
                    if tables and hasattr(tables, 'tables') and tables.tables:
                        table_count = len(tables.tables)
                        total_tables += table_count
                        print(f"页面 {page_idx + 1} 找到 {table_count} 个表格")

                        # 处理每个表格
                        for table_idx, table in enumerate(tables.tables):
                            try:
                                # 打印表格对象信息以辅助调试
                                print(f"表格 {page_idx + 1}-{table_idx + 1} 类型: {type(table)}")
                                if hasattr(table, 'rect'):
                                    print(f"表格 {page_idx + 1}-{table_idx + 1} 区域: {table.rect}")
                                if hasattr(table, 'cells') and table.cells:
                                    print(f"表格 {page_idx + 1}-{table_idx + 1} 单元格数: {len(table.cells)}")

                                # 格式化表格
                                table_text = self._format_table_from_pymupdf(table, page_idx + 1, table_idx + 1)
                                page_text += "\n" + table_text + "\n"
                            except Exception as e:
                                print(f"处理表格 {page_idx + 1}-{table_idx + 1} 时出错: {str(e)}")
                                traceback.print_exc()
                                page_text += f"\n[表格 {page_idx + 1}-{table_idx + 1} 处理失败: {str(e)}]\n"
                    elif hasattr(tables, '__iter__') or (hasattr(tables, '__len__') and len(tables) > 0):
                        # 处理表格列表
                        try:
                            table_list = list(tables)
                            total_tables += len(table_list)
                            print(f"页面 {page_idx + 1} 找到 {len(table_list)} 个表格(列表模式)")

                            for table_idx, table in enumerate(table_list):
                                try:
                                    table_text = self._format_table_from_pymupdf(table, page_idx + 1, table_idx + 1)
                                    page_text += "\n" + table_text + "\n"
                                except Exception as e:
                                    print(f"处理表格列表中的表格 {page_idx + 1}-{table_idx + 1} 时出错: {str(e)}")
                                    traceback.print_exc()
                                    page_text += f"\n[表格 {page_idx + 1}-{table_idx + 1} 处理失败: {str(e)}]\n"
                        except Exception as e:
                            print(f"处理表格列表时出错: {str(e)}")
                            traceback.print_exc()
                    else:
                        # 最后尝试直接处理表格对象
                        try:
                            if tables:
                                print(f"尝试直接处理表格对象")
                                table_text = self._format_table_from_pymupdf(tables, page_idx + 1, 1)
                                page_text += "\n" + table_text + "\n"
                                total_tables += 1
                        except Exception as e:
                            print(f"直接处理表格对象时出错: {str(e)}")
                            traceback.print_exc()
                except Exception as e:
                    print(f"查找页面 {page_idx + 1} 的表格时出错: {str(e)}")
                    traceback.print_exc()

                # 提取图片 - 添加更多错误处理
                try:
                    image_list = page.get_images(full=True)
                    page_images = 0

                    for img_idx, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            if not base_image or "image" not in base_image:
                                print(f"无法提取图片 {page_idx + 1}-{img_idx + 1}")
                                continue

                            image_bytes = base_image["image"]

                            # 生成唯一的临时文件名
                            temp_img_path = f"temp_img_{uuid.uuid4()}_{page_idx}_{img_idx}.png"
                            try:
                                with open(temp_img_path, "wb") as img_file:
                                    img_file.write(image_bytes)

                                # 检查文件是否实际写入
                                if not os.path.exists(temp_img_path) or os.path.getsize(temp_img_path) == 0:
                                    print(f"临时图片文件未写入或为空: {temp_img_path}")
                                    continue

                                # OCR处理图片
                                img_text = extract_text_with_subprocess(temp_img_path)

                                if img_text and len(img_text.strip()) > 10:  # 只保留有意义的OCR结果
                                    page_text += f"\n插图开始\n图片 {page_idx + 1}-{img_idx + 1}:\n{img_text}\n插图结束\n"
                                    page_images += 1
                                    total_images += 1
                            except Exception as e:
                                print(f"处理图片 {page_idx + 1}-{img_idx + 1} 时出错: {str(e)}")
                                traceback.print_exc()
                            finally:
                                # 删除临时文件
                                try:
                                    if os.path.exists(temp_img_path):
                                        os.remove(temp_img_path)
                                except Exception as e:
                                    print(f"无法删除临时图片文件 {temp_img_path}: {str(e)}")
                        except Exception as e:
                            print(f"处理PDF中的图片 {page_idx + 1}-{img_idx + 1} 时出错: {str(e)}")
                            traceback.print_exc()
                except Exception as e:
                    print(f"提取页面 {page_idx + 1} 的图片时出错: {str(e)}")
                    traceback.print_exc()

                # 添加页面文本到完整文本
                full_text += page_text + "\n\n"

            # 关闭文档
            doc.close()

            structure["tables_count"] = total_tables
            structure["images_count"] = total_images

            return full_text, structure

        except Exception as e:
            print(f"处理PDF文件时出现顶层异常: {str(e)}")
            traceback.print_exc()
            # 返回错误信息而不是抛出异常，允许部分处理结果
            return f"PDF处理错误: {str(e)}", structure

    def _format_table_from_pymupdf(self, table, page_idx, table_idx):
        """将PyMuPDF提取的表格格式化为文本"""
        table_content = []
        table_content.append(f"\n表格 {page_idx}-{table_idx}:")
        table_content.append("表格开始")

        try:
            # 检查table对象并提取基本信息
            if table and hasattr(table, 'cells') and table.cells:
                # 打印表格对象的属性以便调试
                print(f"表格属性: {dir(table)}")

                # 尝试多种方式获取表格结构
                # 方式1：使用rect属性分析表格维度
                if hasattr(table, 'rect'):
                    print(f"表格区域: {table.rect}")

                # 方式2：从cells中分析表格结构
                cells_data = []
                cell_rows = {}
                cell_cols = {}

                # 从cells中分析行列结构
                for i, cell in enumerate(table.cells):
                    if hasattr(cell, 'rect'):
                        # 使用单元格位置估计其行列
                        y_pos = cell.rect.y0  # 取单元格顶部y坐标作为行标识
                        x_pos = cell.rect.x0  # 取单元格左侧x坐标作为列标识

                        # 记录唯一的行和列位置
                        cell_rows[y_pos] = cell_rows.get(y_pos, 0) + 1
                        cell_cols[x_pos] = cell_cols.get(x_pos, 0) + 1

                    # 提取单元格文本
                    cell_text = ""
                    if hasattr(cell, 'text'):
                        cell_text = cell.text
                    elif hasattr(table, 'cell_text') and callable(getattr(table, 'cell_text')):
                        try:
                            cell_text = table.cell_text(i)
                        except:
                            cell_text = f"单元格{i}"
                    else:
                        cell_text = f"单元格{i}"

                    # 保存单元格位置和文本
                    if hasattr(cell, 'rect'):
                        cells_data.append((cell.rect.y0, cell.rect.x0, cell_text.strip() if cell_text else ""))

                # 如果成功提取了单元格位置
                if cells_data:
                    # 按行排序单元格
                    cells_data.sort()  # 默认按y0排序

                    # 获取唯一的行位置，按顺序排列
                    rows_pos = sorted(cell_rows.keys())
                    cols_pos = sorted(cell_cols.keys())

                    print(f"检测到表格大小: {len(rows_pos)}行 x {len(cols_pos)}列")

                    # 重建表格结构
                    table_rows = []
                    current_row = []
                    current_row_pos = None

                    for y, x, text in cells_data:
                        # 如果是新的一行
                        if current_row_pos is None or abs(y - current_row_pos) > 5:  # 5是容差值
                            if current_row:
                                # 按x坐标排序单元格
                                current_row.sort(key=lambda item: item[0])
                                table_rows.append([text for _, text in current_row])
                            current_row = [(x, text)]
                            current_row_pos = y
                        else:
                            current_row.append((x, text))

                    # 添加最后一行
                    if current_row:
                        current_row.sort(key=lambda item: item[0])
                        table_rows.append([text for _, text in current_row])

                    # 输出表格内容
                    if table_rows:
                        # 假设第一行是表头
                        headers = table_rows[0] if table_rows else []
                        if headers:
                            table_content.append("表头: " + " | ".join(headers))

                        # 处理表体
                        if len(table_rows) > 1:
                            table_content.append("表体:")
                            for row_idx, row in enumerate(table_rows[1:], 1):
                                row_text = []
                                for col_idx, cell in enumerate(row):
                                    col_name = headers[col_idx] if col_idx < len(headers) and headers[
                                        col_idx] else f"列{col_idx + 1}"
                                    row_text.append(f"{col_name}: {cell}")

                                if row_text:
                                    table_content.append(f"行 {row_idx}: " + " | ".join(row_text))
                    else:
                        # 简单地列出所有单元格
                        table_content.append("表格内容(无法识别结构):")
                        for i, cell in enumerate(table.cells):
                            if hasattr(cell, 'text'):
                                table_content.append(f"单元格 {i + 1}: {cell.text.strip()}")
            else:
                # 回退方案：尝试直接访问表格内容
                table_content.append("使用备用方法提取表格:")

                # 尝试通过dir查看对象属性
                obj_attrs = dir(table)
                print(f"表格对象属性: {obj_attrs}")

                # 检查是否可能是TableFinder对象
                if hasattr(table, 'tables') and table.tables:
                    table_content.append(f"检测到TableFinder对象，包含{len(table.tables)}个表格")
                    return self._format_table_from_pymupdf(table.tables[0], page_idx, table_idx)

                # 通用方法：尝试获取任何可能包含表格数据的属性
                for attr in ['cells', 'spans', 'rows', '_rows', 'data', '_cells', 'content']:
                    if attr in obj_attrs:
                        try:
                            value = getattr(table, attr)
                            if value:
                                table_content.append(f"属性 '{attr}' 内容: {str(value)[:100]}...")
                        except:
                            pass

                table_content.append("表格结构无法识别")

        except Exception as e:
            print(f"格式化表格时出现异常: {str(e)}")
            traceback.print_exc()
            table_content.append(f"表格处理出错: {str(e)}")

        table_content.append("表格结束")
        return "\n".join(table_content)

    def _read_image_file(self, file_path: str) -> str:
        """读取图片文件并使用OCR提取文本"""
        try:
            # 使用extract_text_with_subprocess进行OCR处理
            text = extract_text_with_subprocess(file_path)
            return f"图片内容:\n{text}"
        except Exception as e:
            print(f"OCR处理图片时出错: {str(e)}")
            return "无法识别图片内容"

    def _read_markdown_file(self, file_path: str) -> Tuple[str, Dict]:
        """
        读取Markdown文件并转换为纯文本

        返回:
            Tuple[str, Dict]: 文本内容和文档结构
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()

        # 提取标题结构
        structure = {
            "sections": []
        }

        # 查找所有标题
        header_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        for match in header_pattern.finditer(md_text):
            level = len(match.group(1))
            title = match.group(2)
            position = match.start()

            structure["sections"].append({
                "level": level,
                "title": title,
                "position": position
            })

        # 将Markdown转换为HTML
        html = markdown.markdown(md_text)

        # 使用BeautifulSoup提取纯文本
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text(), structure

    def _read_ppt_file(self, file_path: str) -> Tuple[str, Dict]:
        """
        读取PPT文件，提取文本和图片内容

        返回:
            Tuple[str, Dict]: 文本内容和文档结构
        """
        presentation = pptx.Presentation(file_path)
        content = []
        structure = {
            "slides": len(presentation.slides),
            "has_notes": False,
            "images_count": 0
        }

        total_images = 0
        has_notes = False

        # 创建临时目录保存图片
        temp_dir = "temp_ppt_images"
        os.makedirs(temp_dir, exist_ok=True)

        for slide_num, slide in enumerate(presentation.slides):
            slide_content = [f"\n幻灯片 {slide_num + 1}:"]

            # 提取标题
            if slide.shapes.title:
                slide_content.append(f"标题: {slide.shapes.title.text}")

            # 提取文本内容
            text_content = []
            for shape in slide.shapes:
                if hasattr(shape, "text") and shape.text.strip():
                    text_content.append(shape.text)

            if text_content:
                slide_content.append("文本内容:")
                slide_content.extend(text_content)

            # 提取备注
            if slide.has_notes_slide and slide.notes_slide.notes_text_frame.text.strip():
                has_notes = True
                slide_content.append(f"备注: {slide.notes_slide.notes_text_frame.text.strip()}")

            # 提取图片并进行OCR
            image_count = 0
            for shape in slide.shapes:
                if shape.shape_type == 13:  # MSO_SHAPE_TYPE.PICTURE
                    try:
                        image_count += 1
                        total_images += 1

                        # 保存图片到临时文件
                        img_path = f"{temp_dir}/slide_{slide_num + 1}_img_{image_count}.png"
                        with open(img_path, 'wb') as f:
                            f.write(shape.image.blob)

                        # 使用extract_text_with_subprocess进行OCR处理
                        img_text = extract_text_with_subprocess(img_path)

                        if img_text.strip():
                            slide_content.append("\n插图开始")
                            slide_content.append(f"幻灯片图片 {image_count}:")
                            slide_content.append(f"图片内容: {img_text.strip()}")
                            slide_content.append("插图结束")

                        # 删除临时文件
                        if os.path.exists(img_path):
                            os.remove(img_path)
                    except Exception as e:
                        print(f"处理PPT中的图片时出错: {str(e)}")

            content.append("\n".join(slide_content))

        # 清理临时目录
        if os.path.exists(temp_dir) and len(os.listdir(temp_dir)) == 0:
            os.rmdir(temp_dir)

        structure["has_notes"] = has_notes
        structure["images_count"] = total_images

        return '\n\n'.join(content), structure

    def read_file(self, file_path: str) -> Dict[str, Any]:
        """
        读取文件并提取内容

        参数:
            file_path (str): 文件路径

        返回:
            Dict[str, Any]: 包含文本内容和结构信息的字典
        """
        result = {
            "content": "",
            "structure": {},
            "paragraphs": []
        }

        try:
            # 检查文件是否存在
            print(f"开始读取文件: {file_path}")
            if not os.path.exists(file_path):
                error_msg = f"文件不存在: {file_path}"
                print(error_msg)
                raise FileNotFoundError(error_msg)

            # 检查文件大小
            file_size = os.path.getsize(file_path)
            print(f"文件大小: {file_size} 字节")
            if file_size == 0:
                error_msg = f"文件为空: {file_path}"
                print(error_msg)
                raise ValueError(error_msg)

            # 获取文件扩展名
            _, file_ext = os.path.splitext(file_path.lower())
            print(f"文件扩展名: {file_ext}")

            try:
                if file_ext == '.txt':
                    print(f"处理文本文件...")
                    result["content"] = self._read_text_file(file_path)
                    result["paragraphs"] = self._split_into_paragraphs(result["content"])
                elif file_ext == '.docx':
                    print(f"处理Word文档(docx)...")
                    content, structure = self._read_docx_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                elif file_ext == '.doc':
                    print(f"处理Word文档(doc)...")
                    content, structure = self._read_doc_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                elif file_ext == '.pdf':
                    print(f"处理PDF文件...")
                    content, structure = self._read_pdf_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
                    print(f"处理图像文件...")
                    result["content"] = self._read_image_file(file_path)
                    result["paragraphs"] = self._split_into_paragraphs(result["content"])
                elif file_ext == '.md':
                    print(f"处理Markdown文件...")
                    content, structure = self._read_markdown_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                elif file_ext in ['.ppt', '.pptx']:
                    print(f"处理PowerPoint文件...")
                    content, structure = self._read_ppt_file(file_path)
                    result["content"] = content
                    result["structure"] = structure
                    result["paragraphs"] = self._split_into_paragraphs(content)
                else:
                    # 检查文件类型
                    import mimetypes
                    mime_type, _ = mimetypes.guess_type(file_path)
                    print(f"未知文件类型: {file_ext}, MIME类型: {mime_type}")

                    error_msg = f"不支持的文件类型: {file_ext}"
                    print(error_msg)
                    raise ValueError(error_msg)
            except Exception as e:
                print(f"处理文件时出错: {str(e)}")
                traceback.print_exc()
                raise

            print(f"文件处理成功，提取内容长度: {len(result['content'])}")
            if not result["content"]:
                print(f"警告: 提取的内容为空")

            return result
        except Exception as e:
            print(f"读取文件失败: {str(e)}")
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"读取文件失败: {str(e)}")

